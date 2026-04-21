"""
Build an editable DOCX from OCR'd text files.
Supports RTL (Arabic, Hebrew) and LTR languages with proper formatting.

Usage:
    python build_docx.py --input pages.txt --output output.docx [options]

Options:
    --input         Path(s) to OCR text file(s) with ===PAGE N=== delimiters
    --output        Output DOCX file path
    --font          Font name (default: Sakkal Majalla)
    --font-size     Body text font size in points (default: 16)
    --direction     Text direction: rtl, ltr, or auto (default: auto)
    --title-pages   Number of initial pages to format as title/cover (default: 6)
    --chapter-titles Comma-separated chapter title strings for heading formatting
"""

import argparse
import re
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("Missing dependencies. Install with:")
    print("  pip install python-docx lxml")
    sys.exit(1)


# --- RTL helper functions ---

def set_rtl_paragraph(para):
    """Set RTL direction on paragraph properties."""
    pPr = para._p.get_or_add_pPr()
    etree.SubElement(pPr, qn('w:bidi'))

def set_rtl_run(run):
    """Set RTL direction on run properties."""
    rPr = run._r.get_or_add_rPr()
    etree.SubElement(rPr, qn('w:rtl'))

def set_document_rtl(doc):
    """Set default RTL direction at document level."""
    styles = doc.styles.element
    rPrDefault = styles.find(qn('w:docDefaults'))
    if rPrDefault is None:
        rPrDefault = etree.SubElement(styles, qn('w:docDefaults'))
    rPrDef = rPrDefault.find(qn('w:rPrDefault'))
    if rPrDef is None:
        rPrDef = etree.SubElement(rPrDefault, qn('w:rPrDefault'))
    rPr = rPrDef.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(rPrDef, qn('w:rPr'))
    etree.SubElement(rPr, qn('w:rtl'))

def add_section_rtl(section):
    """Set RTL direction on section properties."""
    sectPr = section._sectPr
    etree.SubElement(sectPr, qn('w:bidi'))


# --- Text parsing ---

def parse_pages(filepath):
    """Parse text file with ===PAGE N=== delimiters into (page_num, lines) tuples."""
    pages = []
    current_lines = []
    current_page = None
    with open(filepath, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            m = re.match(r'===PAGE (\d+)===', line)
            if m:
                if current_page is not None:
                    pages.append((current_page, current_lines))
                current_page = int(m.group(1))
                current_lines = []
            else:
                current_lines.append(line)
    if current_page is not None:
        pages.append((current_page, current_lines))
    return pages


def detect_direction(pages):
    """Auto-detect text direction from content. Returns 'rtl' or 'ltr'."""
    arabic_count = 0
    latin_count = 0
    for _, lines in pages:
        for line in lines:
            for ch in line:
                if '\u0600' <= ch <= '\u06FF' or '\u0750' <= ch <= '\u077F':
                    arabic_count += 1
                elif 'A' <= ch <= 'z':
                    latin_count += 1
    return 'rtl' if arabic_count > latin_count else 'ltr'


# --- DOCX builder ---

def build_docx(pages, output_path, font_name, font_size, direction,
               title_page_count, chapter_titles):
    """Build the DOCX document from parsed pages."""
    is_rtl = direction == 'rtl'
    chapter_set = set(chapter_titles) if chapter_titles else set()
    
    doc = Document()
    
    # Set document direction
    if is_rtl:
        set_document_rtl(doc)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    if is_rtl:
        rPr = style.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = etree.SubElement(style.element, qn('w:rPr'))
        etree.SubElement(rPr, qn('w:rtl'))
    
    # Set margins and section direction
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if is_rtl:
            add_section_rtl(section)
    
    first_page = True
    for page_num, lines in pages:
        # Strip leading/trailing empty lines
        while lines and not lines[0].strip():
            lines = lines[1:]
        while lines and not lines[-1].strip():
            lines = lines[:-1]
        
        if not lines:
            continue
        
        # Page break (except first)
        if not first_page:
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)
        first_page = False
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                para = doc.add_paragraph()
                if is_rtl:
                    set_rtl_paragraph(para)
                continue
            
            # Classify line type
            is_chapter = stripped in chapter_set
            is_separator = stripped == '* * *'
            is_footnote = (
                (stripped.startswith('(') and len(stripped) < 100)
                or stripped.startswith('(*)')
            )
            is_title_page = page_num <= title_page_count
            
            # Create paragraph based on type
            para = doc.add_paragraph()
            
            if is_chapter:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(stripped)
                run.bold = True
                run.font.size = Pt(font_size + 8)
                run.font.name = font_name
            elif is_separator:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run('* * *')
                run.font.name = font_name
            elif is_footnote:
                run = para.add_run(stripped)
                run.font.size = Pt(max(font_size - 4, 8))
                run.font.name = font_name
            elif is_title_page:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(stripped)
                run.font.size = Pt(font_size + 4)
                run.font.name = font_name
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = para.add_run(stripped)
                run.font.name = font_name
                run.font.size = Pt(font_size)
            
            if is_rtl:
                set_rtl_paragraph(para)
                set_rtl_run(run)
    
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"Created: {output_path}")
    print(f"Size: {size_kb:.1f} KB")
    print(f"Pages: {len(pages)}")
    print(f"Font: {font_name} {font_size}pt")
    print(f"Direction: {direction}")


def main():
    parser = argparse.ArgumentParser(
        description='Build editable DOCX from OCR text files'
    )
    parser.add_argument('--input', nargs='+', required=True,
                        help='OCR text file(s) with ===PAGE N=== delimiters')
    parser.add_argument('--output', required=True,
                        help='Output DOCX file path')
    parser.add_argument('--font', default='Sakkal Majalla',
                        help='Font name (default: Sakkal Majalla)')
    parser.add_argument('--font-size', type=int, default=16,
                        help='Body text font size in pt (default: 16)')
    parser.add_argument('--direction', choices=['rtl', 'ltr', 'auto'],
                        default='auto',
                        help='Text direction (default: auto-detect)')
    parser.add_argument('--title-pages', type=int, default=6,
                        help='Number of initial title/cover pages (default: 6)')
    parser.add_argument('--chapter-titles', default='',
                        help='Comma-separated chapter title strings')
    
    args = parser.parse_args()
    
    # Parse all input files
    all_pages = []
    for input_file in args.input:
        all_pages.extend(parse_pages(input_file))
    
    if not all_pages:
        print("Error: No pages found in input files.")
        sys.exit(1)
    
    # Detect direction
    direction = args.direction
    if direction == 'auto':
        direction = detect_direction(all_pages)
        print(f"Auto-detected direction: {direction}")
    
    # Parse chapter titles
    chapter_titles = [t.strip() for t in args.chapter_titles.split(',')
                      if t.strip()] if args.chapter_titles else []
    
    build_docx(
        pages=all_pages,
        output_path=args.output,
        font_name=args.font,
        font_size=args.font_size,
        direction=direction,
        title_page_count=args.title_pages,
        chapter_titles=chapter_titles
    )


if __name__ == '__main__':
    main()
